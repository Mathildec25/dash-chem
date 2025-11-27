"""
Generate Word Report - Pure Python Version
Uses python-docx for document generation
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
from datetime import datetime
import os


def set_cell_background(cell, fill_color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def add_styled_paragraph(doc, text, style_name=None, bold=False, size=None, color=None, alignment=None):
    """Add a paragraph with custom styling"""
    p = doc.add_paragraph()
    if style_name:
        p.style = style_name
    
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment:
        p.alignment = alignment
    
    return p


def create_key_metrics_table(doc, data):
    """Create key metrics summary table"""
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'Value'
    
    for cell in header_cells:
        set_cell_background(cell, '2C3E50')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data rows
    metrics = [
        ('Total Experiments', str(data.get('totalExperiments', 0))),
        ('AI-Suggested Experiments', str(data.get('boExperiments', 0))),
        ('Best Objective Value', data.get('bestResult', {}).get('objectiveValue', 'N/A')),
        ('Improvement', f"{data.get('improvement', 'N/A')}%" if data.get('improvement') else 'N/A')
    ]
    
    for i, (metric, value) in enumerate(metrics, start=1):
        row = table.rows[i]
        row.cells[0].text = metric
        row.cells[1].text = value
        
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return table


def create_best_result_table(doc, best_result, parameters, objectives):
    """Create table showing best experimental result"""
    if not best_result:
        doc.add_paragraph("No data available")
        return
    
    num_rows = 1 + len(parameters or []) + len(objectives or [])
    table = doc.add_table(rows=num_rows, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Parameter/Objective'
    header_cells[1].text = 'Value'
    
    for cell in header_cells:
        set_cell_background(cell, '2C3E50')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    row_idx = 1
    
    # Parameters
    if parameters:
        for param in parameters:
            row = table.rows[row_idx]
            row.cells[0].text = param['name']
            row.cells[1].text = str(best_result.get(param['name'], 'N/A'))
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_idx += 1
    
    # Objectives (with green background)
    if objectives:
        for obj in objectives:
            row = table.rows[row_idx]
            row.cells[0].text = obj['name']
            row.cells[1].text = str(best_result.get(obj['name'], 'N/A'))
            
            for cell in row.cells:
                set_cell_background(cell, 'E8F5E9')
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_idx += 1


def create_top_experiments_table(doc, experiments, parameters, objectives):
    """Create table of top performing experiments"""
    if not experiments or len(experiments) == 0:
        doc.add_paragraph("No experiments available")
        return
    
    num_cols = 1 + len(parameters or []) + len(objectives or [])
    num_rows = 1 + min(len(experiments), 10)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = table.rows[0].cells
    col_idx = 0
    header_cells[col_idx].text = 'Rank'
    col_idx += 1
    
    if parameters:
        for param in parameters:
            header_cells[col_idx].text = param['name']
            col_idx += 1
    
    if objectives:
        for obj in objectives:
            header_cells[col_idx].text = obj['name']
            col_idx += 1
    
    for cell in header_cells:
        set_cell_background(cell, '2C3E50')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data rows
    for exp_idx, exp in enumerate(experiments[:10], start=1):
        row = table.rows[exp_idx]
        col_idx = 0
        
        row.cells[col_idx].text = str(exp_idx)
        col_idx += 1
        
        if parameters:
            for param in parameters:
                row.cells[col_idx].text = str(exp.get(param['name'], 'N/A'))
                col_idx += 1
        
        if objectives:
            for obj in objectives:
                cell = row.cells[col_idx]
                cell.text = str(exp.get(obj['name'], 'N/A'))
                set_cell_background(cell, 'E8F5E9')
                col_idx += 1
        
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_recommendations(doc, data):
    """Add recommendations section"""
    doc.add_heading('Recommendations', level=1)
    
    doc.add_paragraph(
        "Based on the optimization results, the following recommendations are provided:"
    )
    
    # Recommendation 1
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("1. Continue Optimization")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(44, 62, 80)
    
    bo_exp = data.get('boExperiments', 0)
    if bo_exp < 10:
        rec_text = "The optimization is still in early stages. Continue running experiments with the Bayesian Optimization algorithm to further explore the parameter space and identify optimal conditions."
    else:
        rec_text = "The optimization has progressed well. Consider running additional experiments in the most promising regions to fine-tune the optimal parameters."
    
    p = doc.add_paragraph(rec_text)
    p.paragraph_format.left_indent = Inches(0.5)
    
    # Recommendation 2
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("2. Validate Best Results")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(44, 62, 80)
    
    p = doc.add_paragraph(
        "Perform replicate experiments at the best conditions to confirm reproducibility and validate the optimum. "
        "This ensures that the observed performance is reliable and not due to experimental noise."
    )
    p.paragraph_format.left_indent = Inches(0.5)
    
    # Recommendation 3
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("3. Explore Adjacent Regions")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(44, 62, 80)
    
    p = doc.add_paragraph(
        "Consider expanding the search around the best-performing conditions. Small variations in parameters "
        "might reveal even better performance or provide insights into the sensitivity of the process."
    )
    p.paragraph_format.left_indent = Inches(0.5)
    
    # Recommendation 4 for MOBO
    if data.get('optimizationType') == 'MOBO':
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("4. Multi-Objective Trade-offs")
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(44, 62, 80)
        
        p = doc.add_paragraph(
            "Review the Pareto front to understand trade-offs between objectives. Select the solution that best "
            "aligns with your priorities, or consider running additional experiments to better balance conflicting objectives."
        )
        p.paragraph_format.left_indent = Inches(0.5)


def generate_word_report(data, output_path):
    """
    Generate a professional Word report from optimization data
    
    Args:
        data: Dictionary containing report data
        output_path: Path where to save the .docx file
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # ===== TITLE PAGE =====
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("OPTIMIZATION REPORT")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(44, 62, 80)
    
    doc.add_paragraph()
    
    project_name = doc.add_paragraph()
    project_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = project_name.add_run(data.get('projectName', 'Untitled Project'))
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(52, 73, 94)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(127, 140, 141)
    
    type_p = doc.add_paragraph()
    type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = type_p.add_run(f"Optimization Type: {data.get('optimizationType', 'SOBO')}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(127, 140, 141)
    
    doc.add_page_break()
    
    # ===== EXECUTIVE SUMMARY =====
    doc.add_heading('Executive Summary', level=1)
    
    doc.add_paragraph(
        f"This report summarizes the results of a Bayesian Optimization campaign. "
        f"A total of {data.get('totalExperiments', 0)} experiments were conducted, "
        f"including {data.get('boExperiments', 0)} AI-suggested experiments."
    )
    
    doc.add_paragraph()
    doc.add_heading('Key Performance Metrics', level=2)
    create_key_metrics_table(doc, data)
    
    doc.add_paragraph()
    doc.add_heading('Best Experimental Result', level=2)
    create_best_result_table(doc, data.get('bestResult'), data.get('parameters'), data.get('objectives'))
    
    doc.add_page_break()
    
    # ===== DETAILED RESULTS =====
    doc.add_heading('Detailed Results', level=1)
    doc.add_heading('Top Performing Experiments', level=2)
    
    top_n = min(len(data.get('topExperiments', [])), 10)
    doc.add_paragraph(
        f"The following table shows the top {top_n} experiments ranked by objective performance."
    )
    
    create_top_experiments_table(doc, data.get('topExperiments'), data.get('parameters'), data.get('objectives'))
    
    doc.add_page_break()
    
    # ===== VISUALIZATION =====
    if data.get('imagePaths'):
        doc.add_heading('Performance Visualization', level=1)
        
        titles = [
            "Convergence Plot",
            "Objective Distribution",
            "Parallel Coordinates",
            "Parameter Importance"
        ]
        
        for idx, img_path in enumerate(data.get('imagePaths', [])):
            if os.path.exists(img_path):
                doc.add_heading(titles[idx] if idx < len(titles) else f"Plot {idx+1}", level=2)
                try:
                    doc.add_picture(img_path, width=Inches(6))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    doc.add_paragraph(f"[Could not load image: {e}]")
                doc.add_paragraph()
        
        doc.add_page_break()
    
    # ===== SAVE =====
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    import sys
    import json
    
    if len(sys.argv) < 3:
        print("Usage: python generate_report_python.py <data.json> <output.docx>")
        sys.exit(1)
    
    data_path = sys.argv[1]
    output_path = sys.argv[2]
    
    with open(data_path, 'r') as f:
        data = json.load(f)
    
    generate_word_report(data, output_path)
    print(f"Report generated: {output_path}")